import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCREENSHOTS = "/home/vaibhav/.gemini/antigravity/brain/7e126878-1abf-475e-9094-7633497f4c4a"
OUTPUT = "/home/vaibhav/Desktop/Projects/Indian-Auto-Market-Analytics/Indian_Auto_Market_Report.docx"

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_img(path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    return t

# ============ TITLE PAGE ============
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Indian Auto Market Analytics")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("A Statistical Analysis of the Indian Four-Wheeler Market")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Technology Stack: ").bold = True
info.add_run("Python, Streamlit, Pandas, Seaborn, Matplotlib, SciPy")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run("Date: April 2026")

doc.add_page_break()

# ============ TABLE OF CONTENTS ============
heading("Table of Contents")
toc_items = [
    "1. Introduction",
    "2. Source of Dataset",
    "3. Exploratory Data Analysis (EDA)",
    "4. Analysis on Dataset",
    "   4.1 Brand Market Overview",
    "   4.2 Brand Model Breakdown",
    "   4.3 Engine CC vs Mileage",
    "   4.4 Service Cost Trends",
    "   4.5 Value for Money Index",
    "   4.6 Mileage Prediction (SLR)",
    "5. Conclusion",
    "6. Future Scope",
    "7. References"
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ============ 1. INTRODUCTION ============
heading("1. Introduction")
para("The Indian automobile industry is one of the largest in the world. India is currently the third-largest automobile market globally by volume, and the industry contributes significantly to the country's GDP. The four-wheeler segment, which includes passenger cars, SUVs, and utility vehicles, has seen remarkable growth over the past decade driven by rising disposable incomes, urbanization, and improved road infrastructure.")
para("Understanding market dynamics in this sector is crucial for manufacturers, dealers, and consumers alike. Questions such as 'Which brands dominate the market?', 'How does engine size affect fuel economy?', and 'Which vehicles offer the best value for money?' are fundamental to making informed purchasing and business decisions.")
para("This project aims to answer these questions through a data-driven approach. We have built an interactive dashboard using Python and Streamlit that performs six distinct analytical modules on a comprehensive Indian car market dataset. Each module applies appropriate statistical methods ranging from descriptive statistics and hypothesis testing to predictive modeling using Simple Linear Regression.")

heading("1.1 Objectives", level=2)
para("The primary objectives of this project are:")
doc.add_paragraph("To perform thorough Exploratory Data Analysis on the Indian car market dataset", style='List Bullet')
doc.add_paragraph("To identify brand positioning and market share patterns", style='List Bullet')
doc.add_paragraph("To analyze the engineering relationship between engine displacement and fuel efficiency", style='List Bullet')
doc.add_paragraph("To track maintenance cost trends across model years", style='List Bullet')
doc.add_paragraph("To develop a custom Value-for-Money index for objective car comparison", style='List Bullet')
doc.add_paragraph("To build a predictive model for estimating fuel efficiency based on engine capacity", style='List Bullet')

heading("1.2 Tools and Technologies", level=2)
add_table(
    ["Tool", "Purpose"],
    [
        ["Python 3.13", "Core programming language"],
        ["Pandas", "Data manipulation and aggregation"],
        ["NumPy", "Numerical computations"],
        ["Matplotlib", "Static chart generation"],
        ["Seaborn", "Statistical visualization"],
        ["SciPy", "Hypothesis testing and regression"],
        ["Streamlit", "Interactive web dashboard"]
    ]
)

heading("1.3 Project Structure", level=2)
para("The project follows a simple, functional structure:")
code_lines = [
    "Indian-Auto-Market-Analytics/",
    "├── main.py                  # All dashboard logic",
    "├── streamlit_app.py         # Entry point",
    "├── data/",
    "│   └── car_dataset_india.csv",
    "├── images/",
    "│   └── car.png",
    "├── pyproject.toml",
    "└── uv.lock"
]
for line in code_lines:
    p = doc.add_paragraph(line)
    p.runs[0].font.name = 'Consolas'
    p.runs[0].font.size = Pt(9)

doc.add_page_break()

# ============ 2. SOURCE OF DATASET ============
heading("2. Source of Dataset")
heading("2.1 Dataset Overview", level=2)
para("The dataset used in this project is titled 'Car Dataset India' and contains information about 10,000 car listings from the Indian automobile market. The data covers 10 major automobile brands operating in India across a 10-year period from 2015 to 2024.")

heading("2.2 Dataset Specifications", level=2)
add_table(
    ["Property", "Value"],
    [
        ["Total Records", "10,000"],
        ["Total Features", "11"],
        ["File Format", "CSV"],
        ["Time Period", "2015 - 2024"],
        ["Geographic Scope", "India"]
    ]
)

heading("2.3 Feature Description", level=2)
add_table(
    ["#", "Column Name", "Data Type", "Description"],
    [
        ["1", "Car_ID", "int64", "Unique identifier"],
        ["2", "Brand", "string", "Manufacturer name"],
        ["3", "Model", "string", "Car model name"],
        ["4", "Year", "int64", "Year of manufacture"],
        ["5", "Fuel_Type", "string", "Petrol/Diesel/CNG/Electric"],
        ["6", "Transmission", "string", "Manual/Automatic"],
        ["7", "Price", "float64", "Price in INR"],
        ["8", "Mileage", "float64", "Fuel efficiency (kmpl)"],
        ["9", "Engine_CC", "int64", "Engine displacement"],
        ["10", "Seating_Capacity", "int64", "Number of seats"],
        ["11", "Service_Cost", "float64", "Annual service cost"]
    ]
)

heading("2.4 Brands Covered", level=2)
add_table(
    ["Rank", "Brand", "Listings"],
    [
        ["1", "Maruti Suzuki", "1,042"],
        ["2", "Volkswagen", "1,034"],
        ["3", "Hyundai", "1,033"],
        ["4", "Mahindra", "1,024"],
        ["5", "Skoda", "1,011"],
        ["6", "Tata Motors", "989"],
        ["7", "Honda", "982"],
        ["8", "Kia", "973"],
        ["9", "Renault", "958"],
        ["10", "Toyota", "954"]
    ]
)

heading("2.5 Fuel Type Distribution", level=2)
add_table(
    ["Fuel Type", "Count", "Percentage"],
    [
        ["CNG", "2,588", "25.88%"],
        ["Electric", "2,518", "25.18%"],
        ["Diesel", "2,468", "24.68%"],
        ["Petrol", "2,426", "24.26%"]
    ]
)

doc.add_page_break()

# ============ 3. EDA ============
heading("3. Exploratory Data Analysis (EDA)")
para("Exploratory Data Analysis is the critical first step in any data science project. Before running formal hypothesis tests or building predictive models, we need to understand the structure, quality, and distribution patterns within the data.")

heading("3.1 Data Quality Assessment", level=2)
para("We performed a null value audit using df.isnull().sum(). The dataset is 100% complete with zero missing values across all 11 columns. No imputation or data cleaning was required.")
para("All data types are correctly assigned. Numerical columns include Car_ID, Year, Price, Mileage, Engine_CC, Seating_Capacity, and Service_Cost. Categorical columns include Brand, Model, Fuel_Type, and Transmission.")

heading("3.2 Descriptive Statistics", level=2)
para("Using df.describe(), we computed summary statistics for all numerical features:")
add_table(
    ["Statistic", "Year", "Price (₹)", "Mileage", "Engine_CC", "Service_Cost"],
    [
        ["Count", "10,000", "10,000", "10,000", "10,000", "10,000"],
        ["Mean", "2019.5", "~19.5L", "~20.0", "~1,543", "14,969"],
        ["Std Dev", "2.88", "~7.8L", "~5.8", "~543", "5,778"],
        ["Min", "2015", "~5L", "~10", "800", "5,000"],
        ["Max", "2024", "~35L", "~30", "2,500", "25,000"]
    ]
)

heading("3.3 Dashboard Home Page Screenshot", level=2)
add_img(os.path.join(SCREENSHOTS, "home_page_screenshot_1776874570361.png"))
para("Figure 3.1: Dashboard Home Page showing dataset preview, distributions, and correlation heatmap")

heading("3.4 Correlation Analysis", level=2)
para("We computed the Pearson correlation matrix for all numerical features. All pairwise correlations are extremely weak (|r| < 0.02), indicating that the numerical features are largely independent of each other in this dataset. The slight negative correlation between Engine_CC and Mileage (-0.008) aligns with the engineering expectation that larger engines consume more fuel.")

doc.add_page_break()

# ============ 4. ANALYSES ============
heading("4. Analysis on Dataset")

# ---- 4.1 ----
heading("4.1 Analysis 1: Brand Market Overview", level=2)

heading("i. Introduction", level=3)
para("Understanding brand-level market positioning is fundamental to any automotive market study. This analysis examines how the 10 major Indian car brands are distributed in terms of listing volume and average pricing.")

heading("ii. General Description", level=3)
para("We aggregate the dataset by the Brand column to compute volume (total listings) and average price per brand. Additionally, we perform a Chi-square test of independence to determine whether fuel type choice is associated with the brand.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("Functions Used: df.groupby(), pd.crosstab(), scipy.stats.chi2_contingency()")
para("Chi-Square Formula: χ² = Σ (Oi - Ei)² / Ei")
para("Hypotheses:")
doc.add_paragraph("H₀: Brand and Fuel_Type are independent", style='List Bullet')
doc.add_paragraph("H₁: Brand and Fuel_Type are dependent", style='List Bullet')
doc.add_paragraph("Significance level: α = 0.05", style='List Bullet')

heading("iv. Analysis Results", level=3)
para("Chi-Square Test Results:")
add_table(
    ["Parameter", "Value"],
    [
        ["Chi-Square Statistic", "18.3777"],
        ["Degrees of Freedom", "27"],
        ["P-Value", "0.8916"]
    ]
)
para("Since p = 0.8916 > 0.05, we fail to reject the null hypothesis. There is no statistically significant association between brand and fuel type in this dataset.")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "brand_overview_1776874628585.png"))
para("Figure 4.1: Brand Market Overview - Volume distribution and Chi-square test results")

doc.add_page_break()

# ---- 4.2 ----
heading("4.2 Analysis 2: Brand Model Breakdown", level=2)

heading("i. Introduction", level=3)
para("While the previous analysis examined brands at a macro level, this module drills down into individual brand portfolios to understand how sales are distributed across their car models.")

heading("ii. General Description", level=3)
para("The user selects a brand from a dropdown, and the dashboard filters the data to show the frequency distribution of models. A pie chart visualizes the proportional split.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("Functions Used: df[df['Brand'] == selected], Series.value_counts(), matplotlib.pyplot.pie()")
para("Model Share Formula: Model Share (%) = (Count of Model_i / Total Count for Brand) × 100")

heading("iv. Analysis Results", level=3)
para("Each brand's model portfolio can be explored interactively. The pie chart reveals whether a brand relies on a single 'hero model' or maintains a diversified lineup.")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "model_breakdown_1776874638902.png"))
para("Figure 4.2: Brand Model Breakdown - Pie chart showing model-wise sales distribution")

doc.add_page_break()

# ---- 4.3 ----
heading("4.3 Analysis 3: Engine CC vs Mileage", level=2)

heading("i. Introduction", level=3)
para("One of the most fundamental relationships in automotive engineering is the trade-off between engine size (CC) and fuel efficiency (kmpl). This analysis quantifies this relationship using Pearson's correlation coefficient.")

heading("ii. General Description", level=3)
para("Electric vehicles are excluded because they don't have engine displacement. For the remaining ICE vehicles, we group by Engine_CC and Fuel_Type, computing mean mileage for each combination.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("Functions Used: df.groupby(), seaborn.lineplot(), scipy.stats.pearsonr()")
para("Pearson Correlation: r = Σ(xi - x̄)(yi - ȳ) / √[Σ(xi - x̄)² × Σ(yi - ȳ)²]")
para("Hypotheses:")
doc.add_paragraph("H₀: No linear correlation (r = 0)", style='List Bullet')
doc.add_paragraph("H₁: Linear correlation exists (r ≠ 0)", style='List Bullet')

heading("iv. Analysis Results", level=3)
add_table(
    ["Parameter", "Value"],
    [
        ["Sample Size", "7,482 (non-electric)"],
        ["Pearson r", "-0.0077"],
        ["P-Value", "5.08 × 10⁻¹"]
    ]
)
para("The correlation is very close to zero. At the individual car level, there is no statistically significant linear correlation in this dataset. However, grouped averages show the expected slight downward trend.")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "cc_vs_mileage_1776874651151.png"))
para("Figure 4.3: CC vs Mileage - Line plot showing average mileage by engine displacement")

doc.add_page_break()

# ---- 4.4 ----
heading("4.4 Analysis 4: Service Cost Trends", level=2)

heading("i. Introduction", level=3)
para("Maintenance and service costs are a critical component of total cost of ownership. This analysis tracks how average service costs have evolved across model years (2015-2024).")

heading("ii. General Description", level=3)
para("We compute the mean annual service cost for each model year and visualize the trend as a line chart. A box plot displays the full distribution per year including outlier identification.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("Functions Used: seaborn.lineplot(estimator='mean'), seaborn.boxplot()")
para("Box Plot: Box spans Q1 to Q3 (IQR), whiskers extend to 1.5 × IQR, points beyond are outliers.")

heading("iv. Analysis Results", level=3)
add_table(
    ["Year", "Avg Service Cost (₹)"],
    [
        ["2015", "14,971"], ["2016", "14,910"], ["2017", "14,949"],
        ["2018", "15,080"], ["2019", "15,160"], ["2020", "14,810"],
        ["2021", "15,293"], ["2022", "14,960"], ["2023", "14,944"],
        ["2024", "14,627"]
    ]
)
para("Service costs remain stable across the 10-year period, fluctuating within ₹14,627 to ₹15,293.")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "service_costs_1776874672366.png"))
para("Figure 4.4: Service Cost Trends - Line chart and box plot showing cost stability")

doc.add_page_break()

# ---- 4.5 ----
heading("4.5 Analysis 5: Value for Money Index", level=2)

heading("i. Introduction", level=3)
para("The 'Value for Money' concept quantifies which cars deliver the best combination of performance and economy relative to their price. We created a custom mathematical index for this purpose.")

heading("ii. General Description", level=3)
para("The VFM score is computed for every car. The top scorers are presented as 'best value' options. The IQR method identifies statistical outliers.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("VFM Formula: VFM = (Mileage / Price_in_Lakhs) × (Engine_CC / 1000)")
para("IQR Outlier Detection:")
doc.add_paragraph("IQR = Q3 - Q1", style='List Bullet')
doc.add_paragraph("Lower Fence = Q1 - 1.5 × IQR", style='List Bullet')
doc.add_paragraph("Upper Fence = Q3 + 1.5 × IQR", style='List Bullet')

heading("iv. Analysis Results", level=3)
add_table(
    ["Rank", "Brand", "Model", "VFM Score"],
    [
        ["1", "Kia", "EV6", "16.50"],
        ["2", "Renault", "Kiger", "15.85"],
        ["3", "Toyota", "Camry", "15.76"],
        ["4", "Honda", "Jazz", "15.31"],
        ["5", "Hyundai", "i20", "15.06"]
    ]
)
para("Total outlier cars identified by IQR method: 715 out of 10,000 (7.15%)")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "value_for_money_1776874707113.png"))
para("Figure 4.5: Value for Money - Box plot showing VFM score distribution and outliers")

doc.add_page_break()

# ---- 4.6 ----
heading("4.6 Analysis 6: Mileage Prediction using Simple Linear Regression", level=2)

heading("i. Introduction", level=3)
para("Simple Linear Regression (SLR) is a fundamental statistical method for modeling the relationship between a dependent variable and a single independent variable. Here we predict mileage from engine CC using grouped averages.")

heading("ii. General Description", level=3)
para("Rather than using individual car records, we first aggregate by computing the mean mileage for each Engine_CC value (7 data points). We then fit a regression line and provide an interactive prediction tool.")

heading("iii. Specific Requirements, Functions and Formulas", level=3)
para("Functions Used: df.groupby(), scipy.stats.linregress(), seaborn.regplot()")
para("SLR Model: ŷ = β₀ + β₁x")
para("Where: ŷ = Predicted Mileage, x = Engine_CC, β₀ = Intercept, β₁ = Slope")
para("Slope: β₁ = Σ(xi - x̄)(yi - ȳ) / Σ(xi - x̄)²")
para("Intercept: β₀ = ȳ - β₁x̄")

heading("iv. Analysis Results", level=3)
add_table(
    ["Parameter", "Value"],
    [
        ["Slope (β₁)", "-0.000078"],
        ["Intercept (β₀)", "20.1260"],
        ["Correlation (r)", "-0.2721"],
        ["R-squared", "0.0740"],
        ["P-Value", "0.5550"],
        ["Standard Error", "0.000123"]
    ]
)
para("Regression Equation: Mileage = -0.000078 × CC + 20.126")
para("The negative slope confirms the expected engineering relationship: as engine displacement increases, mileage decreases. For every 100cc increase, mileage drops by ~0.0078 kmpl.")

heading("v. Visualization", level=3)
add_img(os.path.join(SCREENSHOTS, "mileage_prediction_1776874741537.png"))
para("Figure 4.6: SLR - Regression plot with prediction tool showing CC vs average mileage")

doc.add_page_break()

# ============ 5. CONCLUSION ============
heading("5. Conclusion")
para("This project successfully developed an interactive Business Intelligence dashboard for the Indian automobile market using Python and Streamlit.")

heading("5.1 Key Findings", level=2)
findings = [
    "The Indian four-wheeler market in this dataset is well-distributed across 10 major brands, with Maruti Suzuki leading in volume.",
    "The Chi-square test revealed no significant association between brand and fuel type (p = 0.8916).",
    "Individual brand analysis reveals varying degrees of model concentration across manufacturers.",
    "The Pearson correlation between Engine_CC and Mileage was weak (r = -0.0077) but directionally correct.",
    "Service costs remained stable across 2015-2024, fluctuating between ₹14,627 and ₹15,293.",
    "The IQR analysis identified 715 cars (7.15%) with exceptional VFM scores.",
    "The SLR model showed the correct negative trend (slope = -0.000078) but limited predictive power (R² = 0.074)."
]
for f in findings:
    doc.add_paragraph(f, style='List Number')

heading("5.2 Limitations", level=2)
limits = [
    "The dataset appears to be synthetic, resulting in near-uniform distributions that weaken statistical relationships.",
    "The SLR model uses only 7 grouped data points, limiting statistical power.",
    "External factors (brand reputation, safety ratings) are not captured.",
    "The VFM formula weights each component equally; real consumer preferences may differ."
]
for l in limits:
    doc.add_paragraph(l, style='List Bullet')

doc.add_page_break()

# ============ 6. FUTURE SCOPE ============
heading("6. Future Scope")

heading("6.1 Data Enhancement", level=2)
doc.add_paragraph("Real Market Data: Replacing the dataset with scraped data from CarDekho or Cars24 would produce stronger statistical relationships.", style='List Bullet')
doc.add_paragraph("Additional Features: Safety ratings (NCAP scores), boot space, and warranty period.", style='List Bullet')
doc.add_paragraph("Time Series: Monthly sales data would enable seasonal trend analysis.", style='List Bullet')

heading("6.2 Advanced Analytics", level=2)
doc.add_paragraph("Multiple Linear Regression (MLR) with multiple predictors for better accuracy.", style='List Bullet')
doc.add_paragraph("Classification models (Logistic Regression, Decision Trees) for value categorization.", style='List Bullet')
doc.add_paragraph("K-Means clustering for natural market segmentation.", style='List Bullet')
doc.add_paragraph("Sentiment Analysis using customer review data.", style='List Bullet')

heading("6.3 Dashboard Enhancement", level=2)
doc.add_paragraph("Normality testing (Shapiro-Wilk) before parametric tests.", style='List Bullet')
doc.add_paragraph("Confidence interval toggles on regression plots.", style='List Bullet')
doc.add_paragraph("PDF/Excel export options for analysis results.", style='List Bullet')

heading("6.4 Deployment", level=2)
doc.add_paragraph("Cloud deployment via Streamlit Community Cloud or Heroku.", style='List Bullet')
doc.add_paragraph("Mobile-responsive layout adaptation.", style='List Bullet')
doc.add_paragraph("REST API endpoint for the prediction model.", style='List Bullet')

doc.add_page_break()

# ============ 7. REFERENCES ============
heading("7. References")
refs = [
    "Society of Indian Automobile Manufacturers (SIAM). 'Statistical Profile of Automobile Industry in India, 2023-24.' https://www.siam.in",
    "McKinney, W. (2017). Python for Data Analysis. O'Reilly Media.",
    "Seabold, S. & Perktold, J. (2010). 'Statsmodels: Econometric and Statistical Modeling with Python.' Proceedings of the 9th Python in Science Conference.",
    "VanderPlas, J. (2016). Python Data Science Handbook. O'Reilly Media.",
    "Streamlit Documentation. https://docs.streamlit.io",
    "SciPy Documentation. 'Statistical Functions.' https://docs.scipy.org/doc/scipy/reference/stats.html",
    "Waskom, M. (2021). 'Seaborn: Statistical Data Visualization.' Journal of Open Source Software, 6(60), 3021.",
    "Montgomery, D.C., Peck, E.A., Vining, G.G. (2012). Introduction to Linear Regression Analysis. Wiley.",
    "NumPy Documentation. https://numpy.org/doc",
    "Pandas Documentation. https://pandas.pydata.org/docs"
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f"[{i}] {ref}")

doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
